from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("/Users/jackychen/Documents/Codex/2026-07-28/1-j/outputs/jacky_chen_task_manager_assessment.docx")

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
TEXT = RGBColor(17, 24, 39)
MUTED = RGBColor(95, 99, 104)
LIGHT = RGBColor(242, 244, 247)


def set_font(run, *, name="Calibri", size=11, color=TEXT, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def add_title(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_font(run, size=23, color=TEXT, bold=True)


def add_subtitle(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run(text)
    set_font(run, size=11, color=MUTED)


def add_heading(doc: Document, text: str, level: int):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    run = paragraph.add_run(text)
    set_font(run, size=16 if level == 1 else 13, color=BLUE if level == 1 else DARK, bold=True)


def add_body(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_font(run)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(item)
        set_font(run)


def add_metadata_table(doc: Document):
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.7), Inches(4.8)]

    rows = [
        ("Candidate", "Jacky Chen"),
        ("Repository", "[Add your GitHub repository URL here]"),
        ("Live Demo", "[Add your deployed frontend URL here]"),
        ("Stack", "React, TypeScript, Vite, Supabase, anonymous auth"),
    ]

    for row_idx, (label, value) in enumerate(rows):
        for col_idx, content in enumerate((label, value)):
            cell = table.rows[row_idx].cells[col_idx]
            cell.width = widths[col_idx]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(content)
            set_font(run, bold=col_idx == 0)
            if col_idx == 0:
                set_cell_shading(cell, "F2F4F7")


def add_schema_table(doc: Document):
    table = doc.add_table(rows=7, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.8), Inches(1.5), Inches(3.2)]
    headers = ["Field", "Type", "Notes"]
    rows = [
      ("id", "uuid", "Primary key"),
      ("title", "text", "Required"),
      ("status", "text", "todo, in_progress, in_review, done"),
      ("description", "text", "Optional detail body, defaults to empty string"),
      ("priority", "text", "low, normal, high"),
      ("user_id", "uuid", "Tied to anonymous guest session via auth.uid()"),
    ]

    for col_idx, label in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.width = widths[col_idx]
        set_cell_shading(cell, "F2F4F7")
        run = cell.paragraphs[0].add_run(label)
        set_font(run, bold=True)

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.width = widths[col_idx]
            run = cell.paragraphs[0].add_run(value)
            set_font(run)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_title(doc, "Task Manager Assessment Submission")
    add_subtitle(
        doc,
        "PulseBoard is a polished Kanban-style task board designed for the internship assessment.",
    )
    add_metadata_table(doc)

    add_heading(doc, "1. Solution Overview", 1)
    add_body(
        doc,
        "PulseBoard is a responsive four-column Kanban board that lets each guest user create tasks, move them across statuses, and visually manage work. The interface emphasizes strong hierarchy, clear spacing, fast scanability, and immediate feedback for loading, empty, and error states.",
    )

    add_heading(doc, "2. Design Decisions", 1)
    add_bullets(
        doc,
        [
            "Used a glassmorphism-inspired dark surface with warm accent gradients to make the board feel intentional instead of generic.",
            "Kept drag-and-drop native with the HTML5 API to reduce dependency risk while preserving direct interaction.",
            "Added summary cards, due date urgency badges, and filtering so the board is more useful than a basic todo list.",
            "Included a demo fallback mode so the product can still be reviewed locally before Supabase credentials are added.",
        ],
    )

    add_heading(doc, "3. Features Implemented", 1)
    add_bullets(
        doc,
        [
            "Anonymous guest session flow through Supabase Auth.",
            "Per-user task isolation with Row Level Security.",
            "Task creation with title, description, priority, and due date.",
            "Drag-and-drop status changes across To Do, In Progress, In Review, and Done.",
            "Search, priority filtering, and due date filtering.",
            "Summary statistics for total, completed, in-flight, overdue, and due-soon tasks.",
            "Realtime-ready subscription hook for task refresh when Supabase mode is enabled.",
        ],
    )

    add_heading(doc, "4. Database Schema", 1)
    add_body(
        doc,
        "The full SQL setup is included in supabase/schema.sql. The main tasks table is summarized below.",
    )
    add_schema_table(doc)
    add_body(
        doc,
        "Additional field included in the actual schema: created_at timestamptz default timezone('utc', now()).",
    )

    add_heading(doc, "5. Setup Instructions", 1)
    add_bullets(
        doc,
        [
            "Install dependencies with pnpm install.",
            "Copy .env.example to .env and set VITE_SUPABASE_URL plus VITE_SUPABASE_ANON_KEY.",
            "Enable Anonymous sign-in in Supabase Auth.",
            "Run the SQL in supabase/schema.sql.",
            "Start locally with pnpm dev.",
            "Deploy the Vite app to Vercel, Netlify, or Cloudflare Pages with the same two environment variables.",
        ],
    )

    add_heading(doc, "6. Tradeoffs and Next Steps", 1)
    add_bullets(
        doc,
        [
            "I prioritized a stable, reviewer-friendly core board over adding a larger optional backend surface.",
            "If I had more time, I would add assignees, task comments, and a dedicated activity log.",
            "I would also add integration tests around anonymous auth bootstrapping and drag-and-drop persistence.",
        ],
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
